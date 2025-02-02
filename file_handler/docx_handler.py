import pytesseract
from PIL import Image
import io
import pandas as pd
import os
import zipfile
from docx import Document

def extract_docx_content(docx_file):
    """
    Extracts text, tables, and images from a DOCX document.
    """
    docx_content = {"text": [], "images": [], "tables": []}
    doc = Document(io.BytesIO(docx_file))
    
    for para in doc.paragraphs:
        if para.text.strip():
            docx_content["text"].append(para.text.strip())
    
    for table in doc.tables:
        table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        docx_content["tables"].append(pd.DataFrame(table_data))
    
    os.makedirs("media", exist_ok=True)
    with zipfile.ZipFile(docx_file, 'r') as docx:
        media_files = [f for f in docx.namelist() if f.startswith('word/media/') and f.lower().endswith(('.jpg', '.png', '.jpeg'))]
        for media in media_files:
            img_bytes = docx.read(media)
            img_extension = media.split('.')[-1]
            img_filename = f"media/{os.path.basename(media)}"
            with open(img_filename, 'wb') as img_file:
                img_file.write(img_bytes)
            docx_content["images"].append(img_filename)
    
    return docx_content
